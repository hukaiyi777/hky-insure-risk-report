# -*- coding: utf-8 -*-
"""养老与传承风险分析报告 · 可复用生成器（暖黄 1.0 版 · 人·权·财 主线）。

用法：
  1) 复制本文件，把顶部 DATA 字典替换为本次客户的内容（结构见 示例先生 示例）。
  2) 家庭架构图通过 DATA['family'] 配置（boxes / lines / labels / note），未确认项 kind='dashed'。
  3) 运行：python report_template.py
     产出：养老与传承风险分析报告.docx  +  family_diagram.png

依赖：pip install python-docx pillow
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================ 暖黄色调（暖黄 1.0 版） ============================
BROWN      = RGBColor(0xA0, 0x6A, 0x30)  # 主色：标题、表格头
DARK_BROWN = RGBColor(0x8B, 0x5A, 0x2B)  # 深色主标题
ORANGE     = RGBColor(0xE0, 0x7A, 0x30)  # 强调色
GREY       = RGBColor(0x8C, 0x7A, 0x6A)  # 暖灰
BODY       = RGBColor(0x4A, 0x35, 0x25)  # 正文暖棕
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

FONT = '微软雅黑'

# ---------------------------- 文档底层工具 ----------------------------
def set_run_font(run, name=FONT, size=None, color=None, bold=None):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name); rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name)
    if size is not None: run.font.size = Pt(size)
    if color is not None: run.font.color.rgb = color
    if bold is not None: run.font.bold = bold

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=10.5, align=None, fill=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    for i, line in enumerate(text.split('\n')):
        if i > 0: p.add_run().add_break()
        run = p.add_run(line); set_run_font(run, size=size, color=color, bold=bold)
    if fill: shade_cell(cell, fill)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)

def add_table(doc, headers, rows, widths=None, header_fill='A06A30', zebra='FFF8F0'):
    ncol = len(headers)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=WHITE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, fill=header_fill)
        set_cell_margins(hdr[i])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            fill = zebra if (ri % 2 == 1) else None
            set_cell(cells[ci], val, size=10.5, fill=fill); set_cell_margins(cells[ci])
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows: r.cells[i].width = Inches(w)
    doc.add_paragraph(); return t

def callout(doc, title, text, fill='FDF3E6', bar='E07A30'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]
    r = p.add_run('◆ ' + title + '\n'); set_run_font(r, size=10.5, color=ORANGE, bold=True)
    for i, line in enumerate(text.split('\n')):
        if i > 0: p.add_run().add_break()
        r2 = p.add_run(line); set_run_font(r2, size=10.5)
    shade_cell(c, fill)
    tcPr = c._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18'); left.set(qn('w:space'), '0'); left.set(qn('w:color'), bar)
    borders.append(left); tcPr.append(borders)
    set_cell_margins(c, top=80, bottom=80, left=120, right=120)
    doc.add_paragraph(); return t

def h(doc, text, size=14, color=BROWN, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_run_font(r, size=size, color=color, bold=True)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6'); bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), 'A06A30')
    pbdr.append(bottom); pPr.append(pbdr); return p

def body(doc, text, size=10.5, color=BODY, bold=False, space=2):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); set_run_font(r, size=size, color=color, bold=bold); return p

def bullet(doc, text, size=10.5, color=BODY):
    p = doc.add_paragraph(style='List Bullet'); r = p.add_run(text)
    set_run_font(r, size=size, color=color); return p

def setup(doc):
    style = doc.styles['Normal']; style.font.name = FONT; style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr(); rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    for s in doc.sections:
        s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)

# ---------------------------- 家庭架构图（Pillow 生成 PNG） ----------------------------
def draw_family_diagram(path, boxes, lines, labels, note):
    """boxes: [{'xy':(x1,y1,x2,y2),'text':...,'kind':'brown'|'orange'|'dashed'}]
    lines: [( (x1,y1),(x2,y2) ), ...]
    labels: [{'xy':(x,y),'text':...,'color':'brown'|'grey'}]
    note: 底部注释字符串
    """
    from PIL import Image, ImageDraw, ImageFont
    W, H = 900, 520
    img = Image.new('RGB', (W, H), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype(r"C:/Windows/Fonts/msyhbd.ttc", 20)
        font_s = ImageFont.truetype(r"C:/Windows/Fonts/msyh.ttc", 16)
        font_xs = ImageFont.truetype(r"C:/Windows/Fonts/msyh.ttc", 14)
    except Exception:
        font = ImageFont.truetype(r"C:/Windows/Fonts/msyh.ttc", 20)
        font_s = ImageFont.truetype(r"C:/Windows/Fonts/msyh.ttc", 16)
        font_xs = ImageFont.truetype(r"C:/Windows/Fonts/msyh.ttc", 14)

    BROWN_RGB = (0xA0, 0x6A, 0x30); ORANGE_RGB = (0xE0, 0x7A, 0x30)
    GREY_RGB = (0x8C, 0x7A, 0x6A); CREAM_RGB = (0xFD, 0xF3, 0xE6)

    def draw_box(d, xy, text, fill=None, outline=None, text_color=(255, 255, 255), dashed=False, f=font):
        x1, y1, x2, y2 = xy
        if fill: d.rectangle(xy, fill=fill)
        if dashed:
            dash, gap = 8, 6
            for x in range(x1, x2, dash + gap):
                d.line([(x, y1), (min(x + dash, x2), y1)], fill=outline, width=2)
                d.line([(x, y2), (min(x + dash, x2), y2)], fill=outline, width=2)
            for y in range(y1, y2, dash + gap):
                d.line([(x1, y), (x1, min(y + dash, y2))], fill=outline, width=2)
                d.line([(x2, y), (x2, min(y + dash, y2))], fill=outline, width=2)
        elif outline: d.rectangle(xy, outline=outline, width=2)
        lines_t = text.split('\n')
        total_h = sum(d.textbbox((0, 0), l, font=f)[3] - d.textbbox((0, 0), l, font=f)[1] for l in lines_t)
        if len(lines_t) > 1: total_h += (len(lines_t) - 1) * 6
        cy = (y1 + y2 - total_h) // 2; cx = (x1 + x2) // 2; y = cy
        for ln in lines_t:
            bb = d.textbbox((0, 0), ln, font=f); tw, th = bb[2] - bb[0], bb[3] - bb[1]
            d.text((cx - tw // 2, y), ln, font=f, fill=text_color); y += th + 6

    def draw_label(d, text, x, y, color=BROWN_RGB, f=font_s):
        bb = d.textbbox((0, 0), text, font=f); tw, th = bb[2] - bb[0], bb[3] - bb[1]
        d.text((x - tw // 2, y - th // 2), text, font=f, fill=color)

    colors = {'brown': BROWN_RGB, 'orange': ORANGE_RGB, 'grey': GREY_RGB}
    for ln in lines:
        draw.line(ln, fill=BROWN_RGB, width=2)
    for lb in labels:
        draw_label(draw, lb['text'], lb['xy'][0], lb['xy'][1], color=colors.get(lb.get('color', 'grey'), GREY_RGB), f=font_xs)
    for b in boxes:
        kind = b.get('kind', 'brown')
        if kind == 'brown':
            draw_box(draw, b['xy'], b['text'], fill=BROWN_RGB, text_color=WHITE)
        elif kind == 'orange':
            draw_box(draw, b['xy'], b['text'], fill=ORANGE_RGB, text_color=WHITE)
        else:  # dashed
            draw_box(draw, b['xy'], b['text'], outline=BROWN_RGB, text_color=BROWN_RGB, dashed=True)

    if note:
        bb = draw.textbbox((0, 0), note, font=font_s); tw, th = bb[2] - bb[0], bb[3] - bb[1]
        nx1, ny1 = (W - tw - 40) // 2, 490; nx2, ny2 = nx1 + tw + 40, ny1 + th + 16
        draw.rounded_rectangle([nx1, ny1, nx2, ny2], radius=6, fill=CREAM_RGB)
        draw.text(((nx1 + nx2 - tw) // 2, (ny1 + ny2 - th) // 2), note, font=font_s, fill=ORANGE_RGB)
    img.save(path, 'PNG')


# ============================ 报告主体 ============================
def build_report(D):
    doc = Document(); setup(doc)

    # ---- 封面 ----
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(70)
    set_run_font(p.add_run(D['client'] + ' · 亲启'), size=22, color=BROWN, bold=True)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_before = Pt(24)
    set_run_font(p2.add_run('养老与传承风险分析报告'), size=28, color=DARK_BROWN, bold=True)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_before = Pt(10)
    # 副标题：交替配色（模板风格）
    sub = D.get('subtitle_parts', [('风险保障', 'BROWN'), (' · ', 'GREY'), ('养老储备', 'ORANGE'),
                                    (' · ', 'GREY'), ('应急保值', 'BROWN'), (' · ', 'GREY'), ('照护与传承安排', 'ORANGE')])
    cmap = {'BROWN': BROWN, 'ORANGE': ORANGE, 'GREY': GREY}
    for txt, col in sub:
        set_run_font(p3.add_run(txt), size=18, color=cmap.get(col, BROWN), bold=True)
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER; p4.paragraph_format.space_before = Pt(36)
    set_run_font(p4.add_run('清流计划 · 守你的财富，护你的晚年，承你的心愿'), size=12, color=BROWN, bold=True)
    for line in D.get('meta_lines', ['顾问：清流计划养老与传承顾问',
                                      '报告性质：风险分析报告（非解决方案）',
                                      '本报告仅供客户本人及授权顾问查阅，未经许可不得外传。',
                                      '解决方案将在二次沟通后，基于确认的信息和客户需求定制。']):
        px = doc.add_paragraph(); px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(px.add_run(line), size=10, color=GREY)
    doc.add_page_break()

    # ---- 服务流程 ----
    h(doc, '我们的服务流程和说明', size=15)
    body(doc, D.get('service_flow_intro', '《养老与传承风险分析报告》是根据多年养老规划经验总结，针对客户家庭或家族的养老风险和需求而设计的一套专业分析服务。'))
    add_table(doc, ['服务步骤', '内容'], D.get('service_flow', [
        ['01 收集信息', '全面评估养老与传承的目标和潜在问题'],
        ['02 分析问题', '分析养老财务目标与潜在问题：人的安排 / 权的安排 / 财的安排'],
        ['03 制定方案', '根据问题制定针对性解决方案：匹配法律工具和金融工具'],
        ['04 回访调整', '定期回顾，适时调整计划，保持长期跟踪'],
    ]), widths=[1.3, 5.5])
    doc.add_page_break()

    # ---- 关于本报告 ----
    h(doc, '一、关于本报告', size=15)
    body(doc, D.get('about_intro', '尊敬的客户：'))
    body(doc, D.get('about_body', '一个真正全面的安排，需要同时兼顾「人、权、财」三大部分：'))
    bullet(doc, '「人」是身心健康、医疗保障与照护安排；')
    bullet(doc, '「权」是意愿的落地执行，如意定监护、遗嘱、传承安排；')
    bullet(doc, '「财」是财富的创造与守护，是让人和权的安排可持续的底气。')
    body(doc, D.get('about_closing', '本报告是一份风险分析报告，而非风险解决方案，目标为全面梳理您在「人·权·财」三条主线上面临的风险。'))
    body(doc, '信息完整度' + D.get('completeness', '约 55%') + '，本报告为初步版；标注「信息确认点」的内容，需在第二次沟通时逐项确认。', color=GREY, size=9.5)
    doc.add_page_break()

    # ---- 客户基本信息确认 ----
    h(doc, '二、客户基本信息确认', size=15)
    body(doc, '下表汇总本次会谈已确认的信息与意愿表达。标注「待确认」的项为信息确认点，将在二次沟通补齐。')
    h(doc, '2.0 家庭架构图（虚线框 = 待确认）', size=12)
    draw_family_diagram('family_diagram.png', D['family']['boxes'], D['family']['lines'], D['family']['labels'], D['family'].get('note', ''))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture('family_diagram.png', width=Inches(5.8)); doc.add_paragraph()

    h(doc, '2.1 基本信息与资产盘点', size=12)
    add_table(doc, ['项目', '已确认信息', '信息来源', '备注'], D['basic_info_rows'], widths=[1.3, 3.0, 1.0, 1.5])
    h(doc, '2.2 资金与意愿表', size=12)
    add_table(doc, ['意愿类型', '客户表达', '风险评估'], D['fund_will_rows'], widths=[1.2, 2.8, 2.8])
    callout(doc, '信息确认点（二次沟通必补）', '\n'.join(D['confirm_points']))
    doc.add_page_break()

    # ---- 多角度风险分析 ----
    h(doc, '三、多角度风险分析', size=15)
    h(doc, '3.1 风险优先级总览', size=12)
    body(doc, '按紧迫度将风险分为 P0（致命且近）至 P3（低/远）四档。')
    add_table(doc, ['等级', '含义', '风险项'], D['risk_overview_rows'], widths=[1.0, 1.5, 4.3])
    h(doc, '3.2 五角度交叉分析', size=12)
    for b in D['five_angle']: bullet(doc, b)
    doc.add_page_break()

    # ---- 人 / 权 / 财 ----
    for key in ['ren', 'quan', 'cai']:
        sec = D[key]
        h(doc, sec['title'], size=15)
        body(doc, sec.get('intro', ''))
        h(doc, sec.get('state_h', '当前状态'), size=12)
        for b in sec.get('state', []): bullet(doc, b)
        h(doc, sec.get('risk_h', '风险点分析'), size=12)
        for b in sec.get('risks', []): bullet(doc, b)
        h(doc, sec.get('confirm_h', '信息确认点'), size=12)
        for b in sec.get('confirm', []): bullet(doc, b)
        doc.add_page_break()

    # ---- 法律释义 ----
    h(doc, '七、法律释义出处', size=15)
    add_table(doc, ['条款 / 文件', '出处', '与本报告的关系'], D['legal_rows'], widths=[1.6, 1.8, 3.4])
    body(doc, '以上法律要点以最新监管与属地机构意见为准，具体落地需协同当地律所与公证处。', color=GREY, size=9.5)

    # ---- 工具矩阵 ----
    h(doc, '八、解决方案预览 / 工具矩阵', size=15)
    body(doc, '以下为市场上可触达的工具方向，只给方向、不推具体产品；具体产品选择、保额设定、品牌推荐将在二次沟通后定制。')
    add_table(doc, ['主线', '可触达工具', '作用'], D['tool_rows'], widths=[1.1, 2.2, 3.5])

    # ---- 风险盘点总表 ----
    h(doc, '九、风险盘点总表', size=15)
    add_table(doc, ['#', '风险项', '紧迫度', '主线', '当前保障', '缺口', '方向'], D['risk_table_rows'], widths=[0.3, 1.6, 0.6, 0.5, 1.3, 1.0, 1.4])

    # ---- 下一步 ----
    h(doc, '十、下一步', size=15)
    h(doc, '10.1 二次沟通议程', size=12)
    for b in D.get('agenda', []): bullet(doc, b)
    h(doc, '10.2 二次沟通前准备的材料', size=12)
    add_table(doc, ['优先级', '材料', '用途'], D['material_rows'], widths=[1.0, 3.0, 2.8])
    doc.add_page_break()

    # ---- 附录 ----
    h(doc, '附录', size=15)
    h(doc, '法律条文要点（摘要）', size=12)
    for b in D.get('legal_summary', []): bullet(doc, b)
    h(doc, '免责声明', size=12)
    body(doc, D.get('disclaimer', '本报告基于一次面谈沟通整理，含多处待确认项，不构成任何保险/投资/法律意见。具体产品配置、保额保费、法律文件起草须在补充信息后，由具备资质的专业人员另行出具方案。'), size=9.5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(20)
    set_run_font(p.add_run('清流计划 · 养老与传承顾问\n守你的财富，护你的晚年，承你的心愿'), size=11, color=BROWN, bold=True)

    doc.save(D.get('out_file', '养老与传承风险分析报告.docx'))
    print('report saved ->', D.get('out_file', '养老与传承风险分析报告.docx'))


# ============================ 示例数据（示例先生） ============================
# 复制本文件后，把下面 DATA 替换为本次客户内容即可。
DATA = {
  'client': '示例先生',
  'completeness': '约 55%',
  'meta_lines': [
      '顾问：清流计划养老与传承顾问',
      '日期：2026年8月7日（首次会谈当日）',
      '报告性质：风险分析报告（非解决方案）',
      '本报告仅供客户本人及授权顾问查阅，未经许可不得外传。',
      '解决方案将在二次沟通后，基于确认的信息和客户需求定制。',
  ],
  'service_flow_intro': '《养老与传承风险分析报告》是根据多年养老规划经验总结，针对客户家庭或家族的养老风险和需求而设计的一套专业分析服务。本报告所有分析都基于您目前的家庭情况、财务状况、生活环境、期望目标，以及我国目前的法律、税务现状和对一些金融参数的合理假设，为您提供一般性的规划建议，不代表我们对实现养老财务目标的保证。',
  'service_flow': [
      ['01 收集信息', '全面评估养老与传承的目标和潜在问题：家庭和财产情况；您对养老与传承的主观想法与期待'],
      ['02 分析问题', '分析养老财务目标与潜在问题：人的安排（健康/医疗/照护）；权的安排（意愿落地/监护/传承）；财的安排（是否充足、是否稳健）'],
      ['03 制定方案', '根据问题制定针对性解决方案：匹配法律工具和金融工具；有针对性地选择产品类型和结构'],
      ['04 回访调整', '定期回顾，适时调整计划，保持长期跟踪'],
  ],
  'about_intro': '尊敬的示例先生：',
  'about_body': '您好，感谢您接受我们的养老与传承规划服务。一个真正全面的安排，需要同时兼顾「人、权、财」三大部分：',
  'about_closing': '本报告是一份风险分析报告，而非风险解决方案。报告的目标是：全面梳理您当前在「人·权·财」三条主线上面临的风险、确认已掌握的信息、标注待核实的关键点。实际解决方案将在完成第二次沟通后，基于确认的完整信息进行定制。',
  'family': {
      'boxes': [
          {'xy': (220, 60, 380, 145), 'text': '母亲\n97岁·已失能\n居敬老院（兄弟分摊）', 'kind': 'brown'},
          {'xy': (520, 60, 680, 145), 'text': '父亲\n[待确认]', 'kind': 'dashed'},
          {'xy': (270, 230, 450, 315), 'text': '示例先生\n61岁·央企退休', 'kind': 'orange'},
          {'xy': (520, 230, 700, 315), 'text': '爱人\n62岁·央企退休', 'kind': 'brown'},
          {'xy': (475, 400, 655, 475), 'text': '独生子\n美国·办身份\n[未婚/回国 待确认]', 'kind': 'dashed'},
      ],
      'lines': [
          ((380, 102), (520, 102)), ((450, 102), (450, 170)), ((450, 170), (360, 170)), ((360, 170), (360, 230)),
          ((450, 272), (520, 272)), ((485, 315), (485, 360)), ((485, 360), (565, 360)), ((565, 360), (565, 400)),
      ],
      'labels': [
          {'xy': (450, 90), 'text': '父母（父亲待确认）', 'color': 'grey'},
          {'xy': (395, 160), 'text': '母子关系', 'color': 'grey'},
          {'xy': (485, 255), 'text': '已婚', 'color': 'grey'},
          {'xy': (600, 360), 'text': '独生子', 'color': 'grey'},
      ],
      'note': '注：长居地未定（海口 / 青岛 / 威海），影响意定监护落地地；父亲情况待确认。',
  },
  'basic_info_rows': [
      ['姓名', '示例先生', '会谈', '—'],
      ['年龄 / 性别', '61岁 / 男', '会谈', '去年央企退休'],
      ['婚姻 / 子女', '已婚，一子（境外）', '会谈', '独生子'],
      ['职业 / 社保', '央企退休职工', '会谈', '社保 + 单位补贴'],
      ['可用资金', '银行储蓄为主，具体金额待确认', '会谈', '海口房挂牌出售中'],
      ['房产', '海口、青岛、威海至少各1套', '会谈', '市值 / 贷款 / 归属待确认'],
      ['父亲情况', '未提及', '—', '待确认'],
      ['母亲情况', '97岁，已失能，居敬老院', '会谈', '由兄弟分摊照护费'],
      ['已有遗嘱', '未提及', '—', '待确认'],
      ['现有保险', '九九鸿福、两全意外险、复星中端医疗、三地惠民保', '会谈', '具体条款 / 现金价值待确认'],
      ['养老目标', '补充退休后保障，重点关注失能/重病/照护', '会谈', '孩子身份与房产出售影响预算'],
  ],
  'fund_will_rows': [
      ['养老目标', '希望到七老八十后仍有保障，尤其失能/重病', '中端医疗不覆盖长护，现金流缺口（P1 · 人）'],
      ['资金分配偏好', '以银行储蓄为主，海口房拟出售', '变现金额与时间不确定，影响可投入预算（P3 · 财）'],
      ['照护偏好', '倾向居家/社区养老，养老社区不强制', '照护资源未采购，失能无执行方案（P2 · 人）'],
      ['传承意愿', '房产/资产给儿子（个人资产）', '儿子境外身份 + 无遗嘱，传承链真空（P1 · 权）'],
      ['时间窗', '等房子出售、孩子身份明朗后再执行', '决策延迟但可接受，服务期两年内可调整（P2 · 财）'],
  ],
  'confirm_points': [
      '①本人及爱人体检报告',
      '②三套房产市值/贷款/产权证明',
      '③银行存款与理财明细',
      '④现有保单清单（两全意外险名称、中端医疗条款页、老保单现金价值）',
      '⑤儿子身份办理进度与时间线',
      '⑥是否已有遗嘱/委托书/医疗预嘱',
      '⑦父亲情况确认',
  ],
  'risk_overview_rows': [
      ['P0（暂无）', '致命且近', '暂无突发断供型风险'],
      ['P1（高）', '影响大、需尽快看见', '失能长护现金流缺口 / 意定监护与医疗预嘱真空 / 跨境传承与资产交接真空'],
      ['P2（中）', '影响中等、可随方案推进', '终身现金流/护理险未配置 / 应急保值账户未独立 / 旧保单低效 / 决策延迟'],
      ['P3（低/远）', '影响小或时间远', '房产处置流动性 / 长护险政策额度有限'],
  ],
  'five_angle': [
      '【时间轴】近期：等房子出售与孩子身份明朗；中期：照护需求上升；远期：失能/传承执行期。',
      '【概率×冲击】集中在低频高损/中频高损：失能长护、境外独子无法决策、跨境继承纠纷。',
      '【可控性】现在能做：补护理险、立遗嘱、定监护人；需提前做：对账、确认身份；只能兜底：监管节奏、寿命通胀。',
      '【关联性】长居地未定→监护人难定；预算未定→护理险规模难定；身份未定→传承结构难定。',
      '【家庭特征】独生子女在境外+双退休+母亲高龄失能=失能无人决策、跨境传承不确定、对长护费用高度敏感。',
  ],
  'ren': {
      'title': '四、人的安排', 'intro': '「人」解决：健康/医疗/照护谁管、钱够不够。',
      'state': ['客户61岁、爱人62岁，自述身体较好，无重大病史（待确认）。',
                '母亲97岁已失能，居敬老院，由兄弟分摊照护费。',
                '已有复星中端医疗 + 三地惠民保 + 单位补充医保。',
                '倾向居家/社区养老，养老社区作为可选项但不强制。'],
      'risks': ['【长期照护现金流缺口】中端医疗像车险只管医院内；养老院日常照护费不报销，高等级约3–4万/月。',
                '【照护资源未锁定】未对接社区日间介护、上门护理、住院护工。',
                '【医疗预嘱空白】无预先医疗指示。',
                '【长护险覆盖有限】2028全国覆盖但基础性、低额度。'],
      'confirm': ['本人及爱人近三年体检报告、既往住院、长期用药、家族病史。',
                  '母亲当前照护费用及兄弟分摊方式。',
                  '对居家/机构养老的真实偏好与可接受费用区间。'],
  },
  'quan': {
      'title': '五、权的安排', 'intro': '「权」解决：意愿怎么落地执行。',
      'state': ['独生子在美国，正办身份，未婚，未来是否回国未定。',
                '长居地未定：海口/青岛/威海。',
                '无遗嘱、委托书、意定监护、医疗预嘱。',
                '境内房产：海口、青岛、威海至少各一套 + 银行储蓄。'],
      'risks': ['【意定监护真空】失能/昏迷时境外独子无法及时决策，极端由法院指定监护人。',
                '【传承意愿落地难】无遗嘱按法定继承，可能引发纠纷。',
                '【跨境继承/外汇/税务不确定】继承结汇、境外税务申报需专业规划。',
                '【监察人未设】缺少监督执行机制。'],
      'confirm': ['是否已有遗嘱/委托书/医疗预嘱。',
                  '儿子身份进度、时间线、回国倾向。',
                  '长居地倾向。',
                  '是否考虑保险金信托、遗赠抚养协议。'],
  },
  'cai': {
      'title': '六、财的安排（三笔钱）', 'intro': '「财」解决：财富的创造与守护，缺三笔钱分层。',
      'state': [],
      'risks': ['【风险保障金】中端医疗+惠民保+两全意外；护理险缺失、重疾/杠杆寿未评估。',
                '【应急储备金】银行储蓄为主、利率下行、海口房变现不确定。',
                '【终身现金流】社保+补贴2万+/月充足；商业年金/增额寿补充未配。'],
      'confirm': ['银行存款/理财/基金/股票明细。',
                  '海口房挂牌价、心理底价、成交时间。',
                  '可配置总预算与年缴上限。',
                  '现有保单条款与现金价值。'],
  },
  'legal_rows': [
      ['第三十三条', '《民法典》总则编', '意定监护制度——预先指定监护人'],
      ['第一千一百二十七条', '《民法典》继承编', '法定继承顺序'],
      ['第一千一百三十三条', '《民法典》继承编', '可通过遗嘱指定继承人/受遗赠人'],
      ['第一千一百四十三条', '《民法典》继承编', '遗嘱的形式要件和有效性'],
      ['《信托法》', '中华人民共和国信托法', '保险金信托——钱不经某人手直接给付'],
      ['《个人外汇管理办法》', '外汇管理法规', '个人年度购汇限额及审批（跨境传承）'],
      ['长护险试点政策', '国家医保局', '2028全国覆盖，基础性、低额度'],
  ],
  'tool_rows': [
      ['人 · 照护', '护理险 + 养老社区/居家养老服务', '护理险触发给付覆盖长护现金流'],
      ['人 · 医疗预嘱', '医疗预嘱 + 意定监护', '提前写明治疗意愿'],
      ['权 · 监护', '意定监护协议 + 监察人', '失能时有人依法决策并受监督'],
      ['权 · 传承', '遗嘱 + 保险金信托 + 跨境结构', '明确归属、处理外汇税务'],
      ['财 · 风险保障', '医疗险检视 + 护理险 + 意外险', '防止大病/失能击穿账户'],
      ['财 · 养老储备', '终身年金保险', '建立活多久领多久现金流'],
      ['财 · 应急保值', '增额终身寿险 + 银行存款', '保持流动性、锁定利率'],
  ],
  'risk_table_rows': [
      ['1', '失能后长期照护现金流缺口', 'P1', '人', '中端医疗/惠民保（不覆盖养老院长护）', '护理险未配', '护理险兜底'],
      ['2', '意定监护与医疗预嘱真空', 'P1', '权', '无', '监护人/预嘱未设', '落地意定监护+预嘱'],
      ['3', '跨境传承与资产交接真空', 'P1', '权', '无遗嘱/信托', '继承规则不明', '遗嘱+传承结构'],
      ['4', '终身现金流/护理险未配置', 'P2', '财', '储蓄为主', '年金/护理险未配', '补充现金流'],
      ['5', '应急保值账户未独立', 'P2', '财', '银行储蓄', '灵活账户未设', '独立应急账户'],
      ['6', '旧保单低效且未审视', 'P2', '财', '九九鸿福保额1万等', '结构错配', '检视/调整'],
      ['7', '房产处置节奏与流动性风险', 'P3', '财', '海口房挂牌', '变现时间不定', '择机处置'],
      ['8', '长护险政策覆盖有限', 'P3', '人/财', '2028全国覆盖（试点）', '额度小', '自有资金+商业险'],
  ],
  'agenda': [
      '逐条确认信息确认点（健康、保单、资金、房产、孩子身份）。',
      '调整风险优先级排序（基于确认后的信息）。',
      '三笔钱的金额锚定——确认预算后框定比例。',
      '讨论传承偏好与意定监护人选。',
  ],
  'material_rows': [
      ['最急', '本人及爱人健康全项（体检/住院/用药/家族病史）', '医疗险与护理险选择基础'],
      ['最急', '现有保单清单（张数、现金价值、受益人、条款页）', '确认保障范围、避免重复'],
      ['次急', '儿子身份办理进度、回国/定居倾向', '传承结构与跨境安排基础'],
      ['次急', '房产市值/贷款/产权证明、存款与理财明细', '框定三笔钱比例'],
      ['可缓', '可配置的年缴预算上限', '护理险/年金规模测算'],
  ],
  'legal_summary': [
      '意定监护：《民法典》第三十三条——可事先书面指定监护人。',
      '遗嘱继承：《民法典》继承编——遗嘱优于法定继承。',
      '信托：《信托法》——保险金信托指定给付、专款专用。',
      '外汇：《个人外汇管理办法》——个人年度购汇限额（跨境传承适用）。',
  ],
  'disclaimer': '本报告基于一次面谈沟通整理，信息完整度约55%，含多处待确认项（信息确认点），不构成任何保险/投资/法律意见，不替代专业核保、法务或税务建议。具体产品配置、保额保费、法律文件起草须在补充信息后，由具备资质的专业人员另行出具方案。',
  'out_file': '养老与传承风险分析报告.docx',
}

if __name__ == '__main__':
    build_report(DATA)
